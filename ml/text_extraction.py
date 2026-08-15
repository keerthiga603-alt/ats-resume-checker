"""
Real text extraction from PDF and DOCX resume files.
No fallback fake text is ever returned; failures raise ExtractionError.
"""
import os
import re


class ExtractionError(Exception):
    pass


def extract_text_from_pdf(filepath: str) -> str:
    import pdfplumber
    if not os.path.exists(filepath):
        raise ExtractionError(f"File not found: {filepath}")
    text_parts = []
    try:
        with pdfplumber.open(filepath) as pdf:
            if len(pdf.pages) == 0:
                raise ExtractionError("PDF has no pages.")
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        raise ExtractionError(f"Unable to extract text from this resume (PDF parse error: {e}).")
    full_text = "\n".join(text_parts).strip()
    if len(full_text) < 30:
        raise ExtractionError(
            "Unable to extract text from this resume. The PDF may be a scanned image "
            "with no selectable text (no OCR layer)."
        )
    return full_text


def extract_text_from_docx(filepath: str) -> str:
    import docx
    if not os.path.exists(filepath):
        raise ExtractionError(f"File not found: {filepath}")
    try:
        d = docx.Document(filepath)
    except Exception as e:
        raise ExtractionError(f"Unable to extract text from this resume (DOCX parse error: {e}).")
    parts = []
    for para in d.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    full_text = "\n".join(parts).strip()
    if len(full_text) < 30:
        raise ExtractionError("Unable to extract text from this resume. The DOCX file appears to be empty.")
    return full_text


def extract_resume_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:
        raise ExtractionError(f"The uploaded file format is not supported ({ext}). Use PDF or DOCX.")


def clean_html_fragments(text: str) -> str:
    """Used when cleaning the training corpus, which contains raw HTML fragments
    (e.g. <span class="hl">...</span>) from the original Indeed scrape."""
    text = re.sub(r"<[^>]+>", " ", text)
    text = text.replace("&nbsp;", " ").replace("&amp;", "&").replace("&#39;", "'")
    text = re.sub(r"\s+", " ", text)
    return text.strip()
